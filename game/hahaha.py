from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


doc = Document()

title = doc.add_heading('🚀 HTML + CSS Beginner Cheat Sheet', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

intro = doc.add_paragraph(
    "A quick reference sheet for building websites while learning web development."
)
intro.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

sections = [
    ("📄 Basic HTML Structure", """<!DOCTYPE html>
<html>
<head>
    <title>My Website</title>
</head>
<body>
    <h1>Hello World</h1>
</body>
</html>"""),
    ("🎨 Background Color", """body {
    background-color: black;
}"""),
    ("🖍️ Text Color", """h1 {
    color: white;
}"""),
    ("📦 Div Example", """<div class="box">
    Content Here
</div>"""),
    ("📐 Width & Height", """width: 300px;
height: 200px;"""),
    ("📏 Padding & Margin", """padding: 20px;
margin: 20px;"""),
    ("🔲 Border Radius", """border-radius: 10px;"""),
    ("✨ Box Shadow", """box-shadow: 0 0 10px gray;"""),
    ("🧲 Flexbox Center", """display: flex;
justify-content: center;
align-items: center;"""),
    ("🖱️ Hover Effect", """button:hover {
    background-color: blue;
}"""),
    ("🔘 Button Example", """button {
    background-color: black;
    color: white;
    padding: 10px 20px;
    border-radius: 10px;
}"""),
    ("🧭 Navbar Example", """.navbar {
    display: flex;
    justify-content: space-between;
    align-items: center;
}"""),
    ("🎯 Useful CSS Properties", """background-color
color
padding
margin
display
justify-content
align-items
font-size
border-radius
box-shadow
transition"""),
    ("💡 Practice Projects", """- Login Page
- Portfolio Website
- Restaurant Website
- To Do App
- Netflix Clone"""),
]

for heading, content in sections:
    doc.add_heading(heading, level=2)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(content)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

doc.add_page_break()

doc.add_heading('🔥 Final Advice', level=2)
doc.add_paragraph(
    "Build while learning. Copy websites, practice daily, and search errors on Google. "
    "The fastest way to improve is by building small projects again and again."
)

path = "HTML_CSS_Cheat_Sheet.docx"
doc.save(path)

print(f"Saved to: {path}")
